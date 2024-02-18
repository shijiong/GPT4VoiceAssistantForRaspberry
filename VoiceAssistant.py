import os
import time  
import azure.cognitiveservices.speech as speechsdk
from openai import AzureOpenAI
import configparser
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime


config = configparser.ConfigParser()
config.read('config.ini')

# Set up OpenAI API credentials
def read_config():


    api_key = config['Azure_OpenAI']['api_key']
    endpoint = config['Azure_OpenAI']['endpoint']
    api_version=config['Azure_OpenAI']['api_version']
    client = AzureOpenAI(azure_endpoint = endpoint,api_key=api_key,api_version=api_version)

    return client

# Set up Azure Speech-to-Text and Text-to-Speech credentials
def select_language():
    print("请选择回答语言：")
    print("1. 普通话")
    print("2. 粤语")
    print("3. 上海话")
    print("4. 四川话")
    print("5. 台湾腔")

    language_choice = input("请输入选项数字: ")

    if language_choice == "1":
        return "zh-CN", "zh-CN", "zh-CN-XiaoxiaoNeural"
    elif language_choice == "2":
        return "yue-CN", "zh-CN", "yue-CN-XiaoMinNeural"
    elif language_choice == "3":
        return "wuu-CN", "zh-CN", "wuu-CN-XiaotongNeural"
    elif language_choice == "4":
        return "zh-CN-sichuan", "zh-CN", "zh-CN-sichuan-YunxiNeural"
    elif language_choice == "5":
        return "zh-TW", "zh-CN", "zh-TW-HsiaoChenNeural"
    else:
        print("无效选项。默认使用普通话。")
        return "zh-CN", "zh-CN", "zh-CN-XiaoxiaoNeural"

speech_synthesis_language, speech_recognition_language, speech_synthesis_voice_name = select_language()

speech_key = config['Azure_Speech_Service']['speech_key']
service_region = config['Azure_Speech_Service']['service_region']
speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)
speech_config.speech_synthesis_language = speech_synthesis_language
speech_config.speech_recognition_language = speech_recognition_language
speech_config.speech_synthesis_voice_name = speech_synthesis_voice_name
speech_synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config)
# speech_key = config['Azure_Speech_Service']['speech_key']
# service_region = config['Azure_Speech_Service']['service_region']
# speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)
# speech_config.speech_synthesis_language = "zh-CN-sichuan" #zh-CN 普通话 wuu-CN 上海话 yue-CN 粤语 zh-CN-sichuan 四川话
# speech_config.speech_recognition_language = "zh-CN"

# speech_config.speech_synthesis_voice_name = "zh-CN-sichuan-YunxiNeural" #zh-CN-XiaoxiaoNeural普通话 zh-TW-HsiaoChenNeural台湾话 wuu-CN-XiaotongNeural上海话 yue-CN-XiaoMinNeural粤语 zh-CN-sichuan-YunxiNeural四川话
# speech_synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config)

# Define the speech-to-text function
def speech_to_text(speech_recognizer):
    # Set up the audio configuration
    audio_config = speechsdk.audio.AudioConfig(use_default_microphone=True)

    # Create a speech recognizer and start the recognition
    result = speech_recognizer.recognize_once_async().get()

    if result.reason == speechsdk.ResultReason.RecognizedSpeech:
        return result.text
    elif result.reason == speechsdk.ResultReason.NoMatch:
        return "I didn't get it, please repeat your question..."
    elif result.reason == speechsdk.ResultReason.Canceled:
        return "Speech recognition canceled."


# Define the text-to-speech function
def text_to_speech(text, rate=1.0):
    try:
        speech_config.speech_synthesis_voice_speed = rate
        result = speech_synthesizer.speak_text_async(text).get()
        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
            #print("Text-to-speech conversion successful.")
            return True
        else:
            print(f"Error synthesizing audio: {result}")
            return False
    except Exception as ex:
        print(f"Error synthesizing audio: {ex}")
        return False



# Define the Azure OpenAI language generation function
def generate_text(prompt, client, context):
    # 将对话历史记录合并为一个文本
    context_text = "\n".join([f"{msg['role']}: {msg['content']}" for msg in context])
    
    # 添加上下文信息到提示文本中
    prompt_with_context = f"{prompt}\nContext:\n{context_text}\n"

    # 使用 OpenAI GPT 生成回复
    response = client.chat.completions.create(
        model="gpt-4-32k",
        messages=[
            {"role": "system", "content": "你是一名人工智能助手，请帮助我们寻找需要的信息."},
            {"role": "user", "content": prompt_with_context}
        ],
    )

    return response.choices[0].message.content

# Define a function to export the conversation to a Word document with different fonts for English and Chinese text
def export_to_word(conversation):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")  # 生成时间戳
    filename = f'conversation_summary_{timestamp}.docx'  # 使用时间戳作为文件名

    doc = Document()
    doc.add_heading('Conversation Summary', level=1)

    # 遍历对话内容，并根据语言设置不同的字体
    for item in conversation:
        if item['role'] == 'system':  # Assistant's response
            p = doc.add_paragraph()
            if any(ord(c) > 127 for c in item['content']):  # Check if the response contains non-ASCII characters
                run = p.add_run("Assistant: " + item['content'])
                run.font.name = '宋体'  # 设置中文宋体字体
                run.font.size = Pt(12)
            else:
                run = p.add_run("Assistant: " + item['content'])
                run.font.name = 'Times New Roman'  # 设置英文Times New Roman字体
                run.font.size = Pt(12)
        elif item['role'] == 'user':  # User's input
            p = doc.add_paragraph()
            if any(ord(c) > 127 for c in item['content']):  # Check if the input contains non-ASCII characters
                run = p.add_run("User: " + item['content'])
                run.font.name = '宋体'  # 设置中文宋体字体
                run.font.size = Pt(12)
            else:
                run = p.add_run("User: " + item['content'])
                run.font.name = 'Times New Roman'  # 设置英文Times New Roman字体
                run.font.size = Pt(12)

    doc.save(filename)
    print(f"Conversation summary exported to {filename}")

# Main program loop
conversation = []  # 用于存储对话内容
export_request = False  # 标志是否收到导出请求
first_iteration=True #使GPT问答更加合理
while True:

    print("我正在倾听，请说“结束”来结束对话。\n如果需要导出到笔记，请说：导出到笔记")
    if first_iteration:
        text_to_speech("我正在倾听，请说“结束”来结束对话。\n如果需要导出到笔记，请说：导出到笔记")
        first_iteration=False
    
    client = read_config()
    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config)
    
    user_input = speech_to_text(speech_recognizer)
    print(f"You said: {user_input}")

    if "结束" in user_input:
        print("Ending the conversation.")
        text_to_speech("感谢您的使用，我们下次再会。")
        break

    if "导出到笔记" in user_input or "请导出笔记" in user_input or "把对话导出" in user_input or "笔记" in user_input:
        export_request = True

    if export_request:
        export_to_word(conversation)
        export_request = False  # 重置导出请求标志
        continue 

    print("Generating response, please wait...")
    prompt = f"Q: {user_input}\nA:"
    response = generate_text(prompt, client,context=conversation)

    print(f"{response}")

    conversation.append({"role": "user", "content": user_input})
    conversation.append({"role": "system", "content": response})

    text_to_speech(response, rate=2.0)

    #目前python docx插件只能修改英文的字形（宋体，仿宋…），中文则一律使用“MS Mincho (中文正文)” 
    #需要手动调整word文档笔记中的字体为宋体